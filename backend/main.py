# -*- coding: utf-8 -*-
from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from typing import List, Optional
from openai import OpenAI
import whisper
import os
from datetime import datetime
import json
from pathlib import Path
import tempfile
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.fonts import addMapping
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
import platform
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="AI Meeting Note API")

# 한글 폰트 등록 함수
def register_korean_fonts():
    """시스템의 한글 폰트를 찾아서 등록"""
    system = platform.system()
    font_registered = False
    
    # Windows 폰트 경로
    if system == "Windows":
        # .env에서 폰트 경로 확인 (우선순위 1)
        custom_font_path = os.getenv("FONT_PATH")
        if custom_font_path and os.path.exists(custom_font_path):
            try:
                pdfmetrics.registerFont(TTFont("Korean", custom_font_path))
                font_registered = True
                print(f"한글 폰트 등록 완료 (사용자 설정): {custom_font_path}")
            except Exception as e:
                print(f"사용자 설정 폰트 등록 실패: {e}")
        
        # 자동 감지 (우선순위 2)
        if not font_registered:
            # Windows 폰트 디렉토리 찾기
            windows_dir = os.environ.get('WINDIR', 'C:/Windows')
            fonts_dir = os.path.join(windows_dir, 'Fonts')
            
            font_paths = [
                os.path.join(fonts_dir, "malgun.ttf"),  # 맑은 고딕
                os.path.join(fonts_dir, "gulim.ttf"),  # 굴림 (TTF)
                os.path.join(fonts_dir, "batang.ttf"),  # 바탕 (TTF)
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont("Korean", font_path))
                        # Bold 폰트도 등록 시도 (맑은 고딕만)
                        if "malgun" in font_path.lower():
                            bold_path = font_path.replace("malgun.ttf", "malgunbd.ttf")
                            if os.path.exists(bold_path):
                                try:
                                    pdfmetrics.registerFont(TTFont("KoreanBold", bold_path))
                                except:
                                    pass
                        font_registered = True
                        print(f"한글 폰트 등록 완료: {font_path}")
                        break
                    except Exception as e:
                        print(f"폰트 등록 실패 {font_path}: {e}")
                        continue
    
    # Mac 폰트 경로
    elif system == "Darwin":
        font_paths = [
            "/System/Library/Fonts/AppleGothic.ttf",
            "/Library/Fonts/AppleGothic.ttf",
            "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
        ]
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("Korean", font_path))
                    font_registered = True
                    print(f"한글 폰트 등록 완료: {font_path}")
                    break
                except Exception as e:
                    print(f"폰트 등록 실패 {font_path}: {e}")
                    continue
    
    # Linux 폰트 경로
    elif system == "Linux":
        font_paths = [
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothicRegular.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        ]
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("Korean", font_path))
                    font_registered = True
                    print(f"한글 폰트 등록 완료: {font_path}")
                    break
                except Exception as e:
                    print(f"폰트 등록 실패 {font_path}: {e}")
                    continue
    
    if not font_registered:
        print("경고: 한글 폰트를 찾을 수 없습니다. PDF에서 한글이 깨질 수 있습니다.")
        print("해결 방법:")
        print("1. Windows: C:/Windows/Fonts/ 폴더에 한글 폰트가 있는지 확인")
        print("2. 폰트가 없다면 맑은 고딕 폰트를 설치하세요")
        print("3. 또는 .env 파일에 FONT_PATH를 설정하세요")
    
    return font_registered

# 앱 시작 시 한글 폰트 등록
KOREAN_FONT_AVAILABLE = register_korean_fonts()

# CORS 설정
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OpenAI 클라이언트 초기화
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Whisper 모델 로드 (한 번만 로드)
# 모델 옵션: tiny, base, small, medium, large
# 성능: tiny < base < small < medium < large
# 정확도: tiny < base < small < medium < large
# 속도: tiny > base > small > medium > large
WHISPER_MODEL_NAME = os.getenv("WHISPER_MODEL", "medium")  # 기본값: medium (base보다 정확함)
whisper_model = None

def get_whisper_model():
    global whisper_model
    if whisper_model is None:
        print(f"Whisper 모델 로딩 중: {WHISPER_MODEL_NAME} (처음 로드 시 시간이 걸릴 수 있습니다)")
        whisper_model = whisper.load_model(WHISPER_MODEL_NAME)
        print(f"Whisper 모델 로드 완료: {WHISPER_MODEL_NAME}")
    return whisper_model

# 파일 저장 디렉토리 설정
MEETINGS_DIR = Path("meetings")
MEETINGS_DIR.mkdir(exist_ok=True)

# 회의록 파일 관리 함수
def get_meeting_file_path(meeting_id: str) -> Path:
    """회의록 파일 경로 반환"""
    return MEETINGS_DIR / f"meeting_{meeting_id}.json"

def save_meeting(meeting_data: dict) -> str:
    """회의록을 JSON 파일로 저장"""
    meeting_id = meeting_data.get("id") or datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    meeting_data["id"] = meeting_id
    file_path = get_meeting_file_path(meeting_id)
    
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(meeting_data, f, ensure_ascii=False, indent=2)
    
    return meeting_id

def load_meeting(meeting_id: str) -> Optional[dict]:
    """회의록 파일 로드"""
    file_path = get_meeting_file_path(meeting_id)
    if not file_path.exists():
        return None
    
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)

def list_meetings() -> List[dict]:
    """모든 회의록 목록 반환"""
    meetings = []
    for file_path in MEETINGS_DIR.glob("meeting_*.json"):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                meeting = json.load(f)
                # 목록 조회용 간단한 정보만 반환
                meetings.append({
                    "id": meeting.get("id", ""),
                    "title": meeting.get("title", ""),
                    "summary": meeting.get("summary", ""),
                    "created_at": meeting.get("created_at", "")
                })
        except Exception as e:
            print(f"파일 로드 오류 {file_path}: {e}")
            continue
    
    # 생성일 기준 내림차순 정렬
    meetings.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return meetings

class SummaryRequest(BaseModel):
    transcript: str
    title: Optional[str] = None

class MeetingResponse(BaseModel):
    id: str
    title: str
    transcript: str
    summary: str
    action_items: List[str]
    created_at: str

@app.get("/")
async def root():
    return {"message": "AI Meeting Note API"}

@app.post("/api/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    save_transcript: bool = Query(False, description="전사 결과를 저장할지 여부")
):
    """음성 파일을 텍스트로 변환
    
    Args:
        file: 오디오 파일
        save_transcript: 전사 결과를 저장할지 여부 (기본값: False)
    """
    try:
        # 임시 파일에 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_path = tmp_file.name
        
        # Whisper로 전사 (성능 향상 옵션 적용)
        model = get_whisper_model()
        result = model.transcribe(
            tmp_path,
            language="ko",  # 한국어 명시
            task="transcribe",  # 전사 작업
            temperature=0.0,  # 더 일관된 결과
            beam_size=5,  # 빔 서치 크기 (정확도 향상)
            best_of=5,  # 최적 결과 선택
            fp16=False,  # CPU 사용 시 False
            verbose=False  # 로그 출력 최소화
        )
        
        transcript = result["text"].strip()
        
        # 세그먼트 정보도 포함 (선택사항)
        segments = []
        if "segments" in result:
            for seg in result["segments"]:
                segments.append({
                    "start": seg.get("start", 0),
                    "end": seg.get("end", 0),
                    "text": seg.get("text", "").strip()
                })
        
        # 전사 결과 저장 (옵션)
        meeting_id = None
        if save_transcript:
            title = f"전사본 {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            meeting_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            
            meeting_data = {
                "id": meeting_id,
                "title": title,
                "transcript": transcript,
                "segments": segments,
                "language": result.get("language", "ko"),
                "summary": "",
                "key_points": [],
                "action_items": [],
                "decisions": [],
                "created_at": datetime.now().isoformat()
            }
            
            save_meeting(meeting_data)
        
        # 임시 파일 삭제
        os.unlink(tmp_path)
        
        return {
            "transcript": transcript,
            "language": result.get("language", "ko"),
            "segments": segments,
            "meeting_id": meeting_id if save_transcript else None
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"전사 오류: {str(e)}")

@app.post("/api/summarize")
async def summarize_meeting(request: SummaryRequest):
    """회의록 요약 및 액션 아이템 추출"""
    try:
        prompt = f"""다음 회의록을 분석하여 다음 형식으로 요약해주세요:

회의록:
{request.transcript}

다음 JSON 형식으로 응답해주세요:
{{
    "summary": "회의의 핵심 내용을 2-3문장으로 요약",
    "key_points": ["주요 논의 사항 1", "주요 논의 사항 2", "주요 논의 사항 3"],
    "action_items": [
        {{"task": "작업 내용", "assignee": "담당자", "deadline": "마감일"}},
        ...
    ],
    "decisions": ["결정 사항 1", "결정 사항 2"]
}}"""

        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional meeting note summarizer. Always respond in Korean with valid JSON format."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        print("response", response)
        summary_text = response.choices[0].message.content
        
        # JSON 파싱 시도
        try:
            summary_json = json.loads(summary_text)
        except:
            # JSON 파싱 실패 시 텍스트 그대로 반환
            summary_json = {
                "summary": summary_text,
                "key_points": [],
                "action_items": [],
                "decisions": []
            }
        
        # 파일로 저장
        title = request.title or f"회의 {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        meeting_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        
        meeting_data = {
            "id": meeting_id,
            "title": title,
            "transcript": request.transcript,
            "summary": summary_json.get("summary", ""),
            "key_points": summary_json.get("key_points", []),
            "action_items": summary_json.get("action_items", []),
            "decisions": summary_json.get("decisions", []),
            "created_at": datetime.now().isoformat(),
            "transcribed_at": datetime.now().isoformat()  # 전사 시간 추가
        }
        
        save_meeting(meeting_data)
        
        return {
            "id": meeting_id,
            "title": title,
            "summary": summary_json.get("summary", ""),
            "key_points": summary_json.get("key_points", []),
            "action_items": summary_json.get("action_items", []),
            "decisions": summary_json.get("decisions", [])
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"요약 오류: {str(e)}")

@app.get("/api/meetings")
async def get_meetings():
    """저장된 회의록 목록 조회"""
    return list_meetings()

@app.get("/api/meetings/{meeting_id}")
async def get_meeting(meeting_id: str):
    """특정 회의록 상세 조회"""
    meeting = load_meeting(meeting_id)
    
    if not meeting:
        raise HTTPException(status_code=404, detail="회의록을 찾을 수 없습니다")
    
    return {
        "id": meeting.get("id", ""),
        "title": meeting.get("title", ""),
        "transcript": meeting.get("transcript", ""),
        "summary": meeting.get("summary", ""),
        "action_items": meeting.get("action_items", []),
        "key_points": meeting.get("key_points", []),
        "decisions": meeting.get("decisions", []),
        "created_at": meeting.get("created_at", "")
    }

@app.get("/api/meetings/{meeting_id}/export/pdf")
async def export_pdf(meeting_id: str):
    """회의록을 PDF로 내보내기"""
    meeting = load_meeting(meeting_id)
    
    if not meeting:
        raise HTTPException(status_code=404, detail="회의록을 찾을 수 없습니다")
    
    # PDF 생성
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        pdf_path = tmp_file.name
    
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    
    # 한글 폰트가 있으면 한글 스타일 생성, 없으면 기본 스타일 사용
    if KOREAN_FONT_AVAILABLE:
        # 한글 지원 스타일 생성
        korean_styles = {
            'Title': ParagraphStyle(
                'KoreanTitle',
                parent=getSampleStyleSheet()['Title'],
                fontName='Korean',
                fontSize=18,
                leading=22,
            ),
            'Heading2': ParagraphStyle(
                'KoreanHeading2',
                parent=getSampleStyleSheet()['Heading2'],
                fontName='Korean',
                fontSize=14,
                leading=18,
            ),
            'Normal': ParagraphStyle(
                'KoreanNormal',
                parent=getSampleStyleSheet()['Normal'],
                fontName='Korean',
                fontSize=11,
                leading=16,
            ),
        }
        styles = korean_styles
    else:
        # 기본 스타일 사용 (한글이 깨질 수 있음)
        styles = getSampleStyleSheet()
        print("경고: 한글 폰트가 없어 기본 폰트를 사용합니다. 한글이 깨질 수 있습니다.")
    
    story = []
    
    # 텍스트 이스케이프 함수 (특수 문자 처리)
    def escape_text(text):
        """HTML 특수 문자 이스케이프"""
        if not text:
            return ""
        text = str(text)
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")
        return text
    
    # 제목
    title_text = escape_text(meeting.get("title", ""))
    story.append(Paragraph(title_text, styles['Title']))
    story.append(Spacer(1, 12))
    
    # 요약
    summary_text = escape_text(meeting.get("summary", "요약 없음"))
    story.append(Paragraph("<b>요약</b>", styles['Heading2']))
    story.append(Paragraph(summary_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # 주요 논의 사항
    key_points = meeting.get("key_points", [])
    if key_points:
        story.append(Paragraph("<b>주요 논의 사항</b>", styles['Heading2']))
        for point in key_points:
            point_text = escape_text(point)
            story.append(Paragraph(f"• {point_text}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # 결정 사항
    decisions = meeting.get("decisions", [])
    if decisions:
        story.append(Paragraph("<b>결정 사항</b>", styles['Heading2']))
        for decision in decisions:
            decision_text = escape_text(decision)
            story.append(Paragraph(f"• {decision_text}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # 액션 아이템
    action_items = meeting.get("action_items", [])
    if action_items:
        story.append(Paragraph("<b>액션 아이템</b>", styles['Heading2']))
        for item in action_items:
            if isinstance(item, dict):
                task = escape_text(item.get("task", ""))
                assignee = escape_text(item.get("assignee", ""))
                deadline = escape_text(item.get("deadline", ""))
                story.append(Paragraph(f"• {task} (담당자: {assignee}, 마감일: {deadline})", styles['Normal']))
            else:
                item_text = escape_text(item)
                story.append(Paragraph(f"• {item_text}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # 전체 전사본
    transcript_text = escape_text(meeting.get("transcript", ""))
    story.append(Paragraph("<b>전체 전사본</b>", styles['Heading2']))
    # 긴 텍스트는 여러 문단으로 나누기
    if transcript_text:
        # 문장 단위로 나누기 (줄바꿈 유지)
        paragraphs = transcript_text.split('\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("전사본이 없습니다.", styles['Normal']))
    
    doc.build(story)
    
    return FileResponse(pdf_path, filename=f"meeting_{meeting_id}.pdf", media_type="application/pdf")

@app.get("/api/meetings/{meeting_id}/export/docx")
async def export_docx(meeting_id: str):
    """회의록을 Word 문서로 내보내기"""
    meeting = load_meeting(meeting_id)
    
    if not meeting:
        raise HTTPException(status_code=404, detail="회의록을 찾을 수 없습니다")
    
    # Word 문서 생성
    doc = Document()
    doc.add_heading(meeting.get("title", ""), 0)
    
    doc.add_heading('요약', level=1)
    doc.add_paragraph(meeting.get("summary", "요약 없음"))
    
    # 주요 논의 사항
    key_points = meeting.get("key_points", [])
    if key_points:
        doc.add_heading('주요 논의 사항', level=1)
        for point in key_points:
            doc.add_paragraph(f"• {point}", style='List Bullet')
    
    # 결정 사항
    decisions = meeting.get("decisions", [])
    if decisions:
        doc.add_heading('결정 사항', level=1)
        for decision in decisions:
            doc.add_paragraph(f"• {decision}", style='List Bullet')
    
    # 액션 아이템
    action_items = meeting.get("action_items", [])
    if action_items:
        doc.add_heading('액션 아이템', level=1)
        for item in action_items:
            if isinstance(item, dict):
                task = item.get("task", "")
                assignee = item.get("assignee", "")
                deadline = item.get("deadline", "")
                doc.add_paragraph(f"• {task} (담당자: {assignee}, 마감일: {deadline})", style='List Bullet')
            else:
                doc.add_paragraph(f"• {item}", style='List Bullet')
    
    doc.add_heading('전체 전사본', level=1)
    doc.add_paragraph(meeting.get("transcript", ""))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        docx_path = tmp_file.name
        doc.save(docx_path)
    
    return FileResponse(docx_path, filename=f"meeting_{meeting_id}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
